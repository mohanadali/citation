import streamlit as st
import pandas as pd
import requests
from pybtex.database import BibliographyData, Entry
from pybtex.database.output.bibtex import Writer
from docx import Document

# Initialize Reference Data
if "references" not in st.session_state:
    st.session_state["references"] = []

# Header
st.title("Enhanced Reference Manager and Citation Generator")

# Upload References
uploaded_file = st.file_uploader("Upload your reference file (BibTeX or RIS)", type=["bib", "ris"])
if uploaded_file:
    file_type = uploaded_file.name.split('.')[-1]
    if file_type == "bib":
        from pybtex.database import parse_string
        bib_data = parse_string(uploaded_file.read().decode("utf-8"), bib_format="bibtex")
        for entry in bib_data.entries.values():
            st.session_state["references"].append({
                "Title": entry.fields.get("title", "N/A"),
                "Author": entry.fields.get("author", "N/A"),
                "Year": entry.fields.get("year", "N/A"),
            })
        st.success("BibTeX file successfully parsed!")

# Display References
if st.session_state["references"]:
    st.write("### Current References")
    references_df = pd.DataFrame(st.session_state["references"])
    st.write(references_df)

    # Edit References
    st.write("### Edit References")
    edited_df = st.experimental_data_editor(references_df, num_rows="dynamic")
    st.session_state["references"] = edited_df.to_dict("records")

# Manual Input for New References
st.write("### Add a New Reference")
with st.form("add_reference"):
    title = st.text_input("Title")
    author = st.text_input("Author")
    year = st.text_input("Year")
    submitted = st.form_submit_button("Add Reference")
    if submitted:
        st.session_state["references"].append({"Title": title, "Author": author, "Year": year})
        st.success("New reference added!")

# Fetch References via DOI (Bulk Input)
st.write("### Fetch References by DOI")
doi_input = st.text_area("Enter DOIs (one per line)")
if st.button("Fetch References"):
    doi_list = doi_input.strip().split("\n")
    for doi in doi_list:
        url = f"https://api.crossref.org/works/{doi.strip()}"
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()["message"]
            new_ref = {
                "Title": data.get("title", ["N/A"])[0],
                "Author": ", ".join([f"{a['given']} {a['family']}" for a in data.get("author", [])]),
                "Year": data.get("published-print", {}).get("date-parts", [[None]])[0][0],
            }
            st.session_state["references"].append(new_ref)
        else:
            st.error(f"Failed to fetch reference for DOI: {doi.strip()}")
    st.success("References fetched and added!")

# Format Citations (With More Styles)
st.write("### Format References")
citation_style = st.selectbox("Choose Citation Style", ["APA", "IEEE", "MLA", "Chicago"])
formatted_references = []
if st.button("Format References"):
    if citation_style == "APA":
        for ref in st.session_state["references"]:
            formatted_references.append(f"{ref['Author']} ({ref['Year']}). {ref['Title']}.")
    elif citation_style == "IEEE":
        for i, ref in enumerate(st.session_state["references"], start=1):
            formatted_references.append(f"[{i}] {ref['Author']}, \"{ref['Title']},\" {ref['Year']}.")
    elif citation_style == "MLA":
        for ref in st.session_state["references"]:
            formatted_references.append(f"{ref['Author']}. \"{ref['Title']}\" ({ref['Year']}).")
    elif citation_style == "Chicago":
        for ref in st.session_state["references"]:
            formatted_references.append(f"{ref['Author']}, \"{ref['Title']}\" ({ref['Year']}).")
    st.write("### Formatted References")
    for ref in formatted_references:
        st.text(ref)

# Export References
st.write("### Export References")
export_format = st.selectbox("Export Format", ["BibTeX", "CSV", "Word (Rich-Text)"])
if st.button("Export"):
    if export_format == "BibTeX":
        bib_entries = {}
        for idx, ref in enumerate(st.session_state["references"], start=1):
            bib_entries[f"ref{idx}"] = Entry(
                "article",
                fields={"title": ref["Title"], "author": ref["Author"], "year": ref["Year"]}
            )
        bib_data = BibliographyData(entries=bib_entries)
        writer = Writer()
        with open("exported_references.bib", "w") as bibfile:
            writer.write_file(bib_data, bibfile)
        st.success("References exported as BibTeX.")
    elif export_format == "CSV":
        export_df = pd.DataFrame(st.session_state["references"])
        export_df.to_csv("exported_references.csv", index=False)
        st.success("References exported as CSV.")
    elif export_format == "Word (Rich-Text)":
        doc = Document()
        doc.add_heading("Formatted References", level=1)
        for ref in formatted_references:
            doc.add_paragraph(ref)
        doc.save("exported_references.docx")
        st.success("References exported as a Word document.")
